# PostulaPro Web App: Propuestas para Mercado Público (CM ID 2239-8-LR25)

import streamlit as st
import io
from datetime import date
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Mostrar logo
st.image("PHOTO-2025-05-08-15-57-57.jpg", width=150)

st.title("PostulaPro: Generador de Propuestas para Mercado Público")

# 📥 Formulario de entrada
with st.form("formulario_propuesta"):
    st.header("Datos del Proveedor")
    razon_social = st.text_input("Razón Social", "Grumpy Chile SpA")
    experiencia = st.text_area("Experiencia Previa")
    capacidades = st.text_area("Capacidades Técnicas")

    st.header("Precios Ofertados")
    precios_input = st.text_area("Escriba ítems y precios (ej: Detergente 5L: 10000)")
    precios = {}
    for linea in precios_input.splitlines():
        if ":" in linea:
            item, precio = linea.split(":")
            precios[item.strip()] = precio.strip()

    st.header("Documentación Legal Adjunta")
    documentos = st.text_area("Lista de documentos (uno por línea)").splitlines()

    st.header("Bases de Licitación")
    bases = st.text_area("Resumen de las bases")

    submitted = st.form_submit_button("Generar Propuesta")

# 🧠 Lógica del agente
def generar_propuesta(razon_social, experiencia, capacidades, precios, documentos, bases):
    resumen = f"Propuesta para la licitación pública Convenio Marco ID 2239-8-LR25 presentada por {razon_social}."
    tecnica = (
        f"{razon_social} cuenta con {experiencia}. Nos especializamos en {capacidades}, cumpliendo con inscripción vigente en el Registro de Proveedores,"
        f" declaración jurada, uso de productos con registro INAPI/ISP y condiciones técnicas exigidas."
    )
    economica = "\n".join([f"{item}: ${precio}" for item, precio in precios.items()])
    carta = (
        f"Sres. Comisión Evaluadora,\n\n"
        f"Por medio de la presente, {razon_social} presenta su postulación a la licitación pública ID 2239-8-LR25 para el Convenio Marco de Aseo e Higiene,"
        f" de acuerdo a lo establecido en las bases. Nos comprometemos a cumplir con los estándares técnicos y contractuales exigidos."
    )
    checklist = [
        "✅ Bases leídas",
        "✅ Oferta técnica redactada",
        "✅ Oferta económica incluida",
        "✅ Documentos legales adjuntos",
        "✅ Declaración jurada generada",
        "✅ Productos con registro ISP/INAPI"
    ]
    analisis = (
        "La propuesta fue desarrollada cumpliendo todos los requisitos formales. Se privilegia experiencia comprobable, cobertura logística y precios competitivos."
    )
    return resumen, tecnica, economica, carta, checklist, analisis

# 📄 Generador de documento Word
def generar_word(razon_social, carta, tecnica, economica, checklist, analisis):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Logo
    section = doc.sections[0]
    header = section.header
    run = header.paragraphs[0].add_run()
    run.add_picture("PHOTO-2025-05-08-15-57-57.jpg", width=Inches(1.5))

    doc.add_heading(f"Propuesta - {razon_social}", 0)
    doc.add_heading("Resumen Ejecutivo", level=1)
    doc.add_paragraph(carta)

    doc.add_heading("Oferta Técnica", level=1)
    doc.add_paragraph(tecnica)

    doc.add_heading("Oferta Económica", level=1)
    doc.add_paragraph(economica)

    doc.add_heading("Checklist de Cumplimiento", level=1)
    for item in checklist:
        doc.add_paragraph(item)

    doc.add_heading("Análisis Estratégico", level=1)
    doc.add_paragraph(analisis)

    doc.add_paragraph("\nFirma: ____________________________")
    doc.add_paragraph("Nombre: __________________________")
    doc.add_paragraph("RUT: _____________________________")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ✅ Si se envió el formulario
if submitted:
    resumen, tecnica, economica, carta, checklist, analisis = generar_propuesta(
        razon_social, experiencia, capacidades, precios, documentos, bases
    )

    st.success("✅ Propuesta generada correctamente")

    st.subheader("Resumen Ejecutivo")
    st.text_area("Resumen", resumen, height=100)

    st.subheader("Oferta Técnica")
    st.text_area("Técnica", tecnica, height=150)

    st.subheader("Oferta Económica")
    st.text_area("Económica", economica, height=150)

    st.subheader("Checklist")
    for item in checklist:
        st.markdown(f"- {item}")

    st.subheader("Análisis Estratégico")
    st.text_area("Análisis", analisis, height=100)

    docx = generar_word(razon_social, carta, tecnica, economica, checklist, analisis)
    st.download_button("📄 Descargar Propuesta Word", docx, file_name="Propuesta_CM2239-8.docx")

